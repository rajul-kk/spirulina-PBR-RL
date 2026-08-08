"""
format_docx.py
Post-processes pandoc-generated .docx files to apply:
  - Plain colourless tables via Word's built-in "Table Grid" style
    (header row text made bold; no fills or colour)
  - All heading levels bold, black, controlled sizes
  - Proper paragraph/heading spacing
  - Consistent Calibri body font

Usage:
    python format_docx.py
    (close any open Word windows for the output files first)
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Font settings ─────────────────────────────────────────────────────────────
BODY_FONT  = "Times New Roman"
BODY_SIZE  = Pt(12)
TABLE_FONT = "Times New Roman"
TABLE_SIZE = Pt(12)
CODE_FONT  = "Courier New"
CODE_SIZE  = Pt(10)
HEADING_FONT = "Arial"  # Separate font for headings

# ── Heading config ────────────────────────────────────────────────────────────
HEADING_SIZES = {
    "Heading 1": 20,
    "Heading 2": 17,
    "Heading 3": 15,
    "Heading 4": 12,
}
HEADING_SPACE_BEFORE = {
    "Heading 1": 18,
    "Heading 2": 14,
    "Heading 3": 10,
    "Heading 4": 8,
}
HEADING_SPACE_AFTER = 6   # pt, all levels


# ── XML helpers ───────────────────────────────────────────────────────────────

def _get_or_add(parent, tag):
    el = parent.find(qn(tag))
    if el is None:
        el = OxmlElement(tag)
        parent.append(el)
    return el


def _set_para_spacing(para, space_before_pt, space_after_pt, line_pt=None):
    pPr = _get_or_add(para._p, "w:pPr")
    spc = _get_or_add(pPr, "w:spacing")
    spc.set(qn("w:before"), str(int(space_before_pt * 20)))
    spc.set(qn("w:after"),  str(int(space_after_pt  * 20)))
    if line_pt is not None:
        spc.set(qn("w:line"),     str(int(line_pt * 20)))
        spc.set(qn("w:lineRule"), "exact")


def _clear_cell_shading(cell):
    """Remove any existing fill so the cell is plain white."""
    tc   = cell._tc
    tcPr = _get_or_add(tc, "w:tcPr")
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "FFFFFF")
    tcPr.append(shd)


# ── Table formatter ───────────────────────────────────────────────────────────

def _force_table_borders(tbl):
    """
    Remove any existing w:tblBorders inserted by pandoc, then inject
    explicit 0.5 pt (sz=4) single black borders for all six sides.
    This overrides pandoc's direct XML so the Table Grid style is not silently
    suppressed.
    """
    tblPr = _get_or_add(tbl, "w:tblPr")

    # Strip whatever pandoc wrote
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)

    # Inject fresh borders: 0.5 pt = sz '4' in OOXML eighths-of-a-point
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")        # 0.5 pt
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")   # black
        tblBorders.append(el)
    tblPr.append(tblBorders)


def format_table(table, doc):
    """
    Apply Word's built-in 'Table Grid' style then force explicit borders.
    Header row text is bold. All cells use TABLE_FONT black.
    Handles <br> tags by inserting actual Word line breaks.
    """
    try:
        table.style = doc.styles["Table Grid"]
    except KeyError:
        pass

    tbl   = table._tbl
    tblPr = _get_or_add(tbl, "w:tblPr")
    tblW = _get_or_add(tblPr, "w:tblW")
    tblW.set(qn("w:type"), "pct")
    tblW.set(qn("w:w"),    "5000")

    _force_table_borders(tbl)

    for row_idx, row in enumerate(table.rows):
        is_header = (row_idx == 0)
        for cell in row.cells:
            _clear_cell_shading(cell)
            for para in cell.paragraphs:
                _set_para_spacing(para, space_before_pt=2, space_after_pt=2)
                
                # Robustly handle line breaks [BR] in all text (including hyperlinks)
                for t in para._p.xpath('.//w:t'):
                    if t.text and "[BR]" in t.text:
                        parts = t.text.split("[BR]")
                        t.text = parts[0]
                        parent_r = t.getparent()
                        # Insert breaks and remaining text into the same run
                        for part in parts[1:]:
                            br = OxmlElement("w:br")
                            parent_r.append(br)
                            new_t = OxmlElement("w:t")
                            new_t.text = part
                            # Ensure we preserve leading/trailing spaces if any
                            new_t.set(qn("xml:space"), "preserve")
                            parent_r.append(new_t)

                # Apply font to all runs (including those inside hyperlinks)
                for r in para._p.xpath('.//w:r'):
                    rPr = _get_or_add(r, "w:rPr")
                    rFonts = _get_or_add(rPr, "w:rFonts")
                    rFonts.set(qn("w:ascii"), TABLE_FONT)
                    rFonts.set(qn("w:hAnsi"), TABLE_FONT)
                    
                    sz = _get_or_add(rPr, "w:sz")
                    sz.set(qn("w:val"), str(int(TABLE_SIZE.pt * 2)))
                    
                    color = _get_or_add(rPr, "w:color")
                    color.set(qn("w:val"), "000000")
                    
                    if is_header:
                        b = _get_or_add(rPr, "w:b")
                        b.set(qn("w:val"), "1")


# ── Heading formatter ─────────────────────────────────────────────────────────

def format_headings(doc):
    """All headings: bold, black, controlled pt sizes, proper spacing."""
    for style_name, size in HEADING_SIZES.items():
        try:
            style            = doc.styles[style_name]
            style.font.name  = HEADING_FONT
            style.font.size  = Pt(size)
            style.font.bold  = True
            style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        except KeyError:
            pass

    for para in doc.paragraphs:
        if para.style.name not in HEADING_SIZES:
            continue
        size = HEADING_SIZES[para.style.name]
        sb   = HEADING_SPACE_BEFORE.get(para.style.name, 12)
        _set_para_spacing(para,
                          space_before_pt=sb,
                          space_after_pt=HEADING_SPACE_AFTER)
        for run in para.runs:
            run.font.name      = HEADING_FONT
            run.font.size      = Pt(size)
            run.font.bold      = True
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


# ── Body font ─────────────────────────────────────────────────────────────────

def format_body(doc):
    """Calibri 11pt body; preserve code blocks; 6pt spacing after paragraphs."""
    code_style_names = {"Source Code", "Verbatim", "Verbatim Char",
                        "Code", "Code Block"}
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            continue
        if para.style.name in code_style_names:
            for run in para.runs:
                run.font.name = CODE_FONT
                run.font.size = CODE_SIZE
            continue
        _set_para_spacing(para, space_before_pt=0, space_after_pt=6,
                          line_pt=13.2)
        for run in para.runs:
            if run.font.name in ("Courier New", "Consolas"):
                run.font.size = CODE_SIZE
                continue
            run.font.name = BODY_FONT
            run.font.size = BODY_SIZE


# ── Page numbers ──────────────────────────────────────────────────────────────

def add_page_numbers(doc):
    """Inserts a dynamic PAGE field in bottom-right footer for each section."""
    for section in doc.sections:
        footer = section.footer
        # If the footer has no paragraphs, create one
        if not footer.paragraphs:
            para = footer.add_paragraph()
        else:
            para = footer.paragraphs[0]
        
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Clear existing runs just in case
        for r in para.runs:
            r.clear()
            
        run = para.add_run()
        run.font.name = BODY_FONT
        run.font.size = Pt(10)
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


# ── Main pipeline ─────────────────────────────────────────────────────────────

def process(input_path, output_path):
    print(f"Processing {input_path} -> {output_path}")
    doc = Document(input_path)

    format_headings(doc)
    format_body(doc)
    add_page_numbers(doc)
    for table in doc.tables:
        format_table(table, doc)

    doc.save(output_path)
    print(f"  OK: Saved ({output_path})")


if __name__ == "__main__":
    targets = [
        ("genetic_env_raw.docx",  "genetic_env_formatted.docx"),
        ("light_env_raw.docx",    "light_env_formatted.docx"),
    ]
    for src, dst in targets:
        process(src, dst)
    print("\nDone.")
